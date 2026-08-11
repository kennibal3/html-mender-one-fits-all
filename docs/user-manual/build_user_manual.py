#!/usr/bin/env python3
"""Build the V1.0 user manual DOCX from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/user-manual/HTML课件互动编辑系统-V1.0-用户手册.md"
OUTPUT = ROOT / "docs/user-manual/output/HTML课件互动编辑系统-V1.0-用户手册.docx"
LOGO = ROOT / "design/brand/preview/logo-lockup-zh-browser.png"

BLUE = "123B5D"
TEAL = "008B7A"
INK = "17211D"
MUTED = "5F6D68"
PALE = "E9F5F2"
WARM = "FAF7F0"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, east_asia: str = "Arial Unicode MS", latin: str = "Aptos") -> None:
    # LibreOffice on macOS resolves the ASCII font first even for mixed CJK
    # runs, so use the installed CJK family for every OOXML font slot.
    run.font.name = east_asia
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        run._element.rPr.rFonts.set(qn(f"w:{slot}"), east_asia)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" 页")
    set_run_font(run)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    # Named override for the compact_reference_guide preset: A4 is used for
    # Chinese software registration and office printing.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{slot}"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, color, before, after in [
        ("Title", 28, BLUE, 0, 12),
        ("Subtitle", 15, TEAL, 0, 10),
        ("Heading 1", 18, BLUE, 14, 8),
        ("Heading 2", 13, TEAL, 10, 5),
        ("Heading 3", 11, BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{slot}"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Figure Caption"]
    caption.font.name = "Arial Unicode MS"
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        caption._element.rPr.rFonts.set(qn(f"w:{slot}"), "Arial Unicode MS")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{slot}"), "Arial Unicode MS")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.65)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(3)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_cell_width(table.cell(0, 0), 7500)
    set_cell_width(table.cell(0, 1), 2100)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("HTML课件互动编辑系统")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(r)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("V1.0 用户手册")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    set_run_font(r)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(1.2)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Cm(10.8))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", "HTML课件互动编辑系统标志")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.0)
    r = p.add_run("HTML课件互动编辑系统")
    set_run_font(r)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("V1.0 用户手册")
    set_run_font(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(1.1)
    r = p.add_run("面向教师的本地 HTML 课件可视化编辑与互动制作工具")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(r)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    labels = ["软件版本", "文档版本", "发布日期", "适用平台", "著作权人"]
    values = ["V1.0", "1.0", "2026年8月", "macOS / Windows", "____________________"]
    for i, (label, value) in enumerate(zip(labels, values)):
        set_cell_width(table.cell(i, 0), 2600)
        set_cell_width(table.cell(i, 1), 5200)
        table.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.cell(i, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), PALE)
        for cell, text_value, bold in [
            (table.cell(i, 0), label, True),
            (table.cell(i, 1), value, False),
        ]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text_value)
            r.bold = bold
            r.font.size = Pt(10)
            set_run_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(2.2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本手册界面截图均来自 V1.0 实际运行环境")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(r)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("目录", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "打开 Word 后右键更新目录")
    p = doc.add_paragraph()
    r = p.add_run("提示：在 Microsoft Word 中打开文档后，可右键目录并选择“更新域”，生成或刷新页码。")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(r)
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1（\2）", text)
    return text.replace("  ", " ").strip()


def add_paragraph_text(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    for idx, part in enumerate(re.split(r"(`[^`]+`)", text)):
        if not part:
            continue
        code = part.startswith("`") and part.endswith("`")
        value = part[1:-1] if code else clean_inline(part)
        r = p.add_run(value)
        set_run_font(r, east_asia="Arial Unicode MS", latin="Aptos Mono" if code else "Aptos")
        if code:
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor.from_string(BLUE)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(clean_inline(text))
    r.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(r)


def add_image(doc: Document, alt: str, relative_path: str) -> None:
    image_path = SOURCE.parent / relative_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Cm(16.4))
    inline_shape._inline.docPr.set("descr", alt)
    inline_shape._inline.docPr.set("title", alt.split("：", 1)[0])
    caption = doc.add_paragraph(alt, style="Figure Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable_twips = 9640
    width = usable_twips // cols
    for row_idx, values in enumerate(rows):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(clean_inline(values[col_idx] if col_idx < len(values) else ""))
            r.font.size = Pt(9)
            r.bold = row_idx == 0
            r.font.color.rgb = RGBColor.from_string(WHITE if row_idx == 0 else INK)
            set_run_font(r)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    block = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        values = [part.strip() for part in lines[idx].strip().strip("|").split("|")]
        block.append(values)
        idx += 1
    if len(block) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in block[1]):
        block.pop(1)
    return block, idx


def add_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## 文档说明")
    idx = start
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            idx += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:]
            p = doc.add_paragraph(title, style="Heading 1")
            p.paragraph_format.keep_with_next = True
            idx += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            p.paragraph_format.keep_with_next = True
            idx += 1
            continue
        if stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:], style="Heading 3")
            idx += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            idx += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered:
            add_paragraph_text(doc, ordered.group(2), "List Number")
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            value = bullet.group(1)
            if value.startswith("[ ] "):
                value = "☐ " + value[4:]
            add_paragraph_text(doc, value, "List Bullet")
            idx += 1
            continue

        paragraph_parts = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if not nxt or re.match(r"^(#{2,4}\s|[-|>]\s?|\d+\.\s|!\[)", nxt):
                break
            paragraph_parts.append(nxt)
            idx += 1
        add_paragraph_text(doc, " ".join(paragraph_parts))


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "HTML课件互动编辑系统 V1.0 用户手册"
    props.subject = "安装、课件编辑、互动制作、版本管理与导出说明"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "HTML课件, 互动编辑, 用户手册, V1.0"
    props.comments = ""


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    set_doc_defaults(doc)
    set_document_properties(doc)
    add_cover(doc)
    add_toc(doc)
    add_body(doc, markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
